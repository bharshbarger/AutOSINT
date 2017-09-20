#!/usr/bin/env python
"""module to generate a docx report based on osint findings"""

#https://python-docx.readthedocs.io/en/latest/user/text.html
#https://python-docx.readthedocs.io/en/latest/user/quickstart.html
import time
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn



class Reportgen(object):
    """class to generate a docx report from module output"""
    def __init__(self):
        self.today = time.strftime("%m/%d/%Y")

    def run(self, \
        args, \
        report_directory, \
        lookup, whois_result, \
        dns_result, \
        google_result, \
        shodan_result, \
        paste_scrape_result, \
        theharvester_result, \
        webscrape_result, \
        cred_result, \
        pyfoca_result):

        for l in lookup:

            print('[+] Starting OSINT report for '.format(l))

            self.document = docx.Document()
            #add logo
            self.document.add_picture('./resources/logo.png', height=Inches(1.25))

            #add domain cover info
            paragraph = self.document.add_paragraph()
            run_paragraph = paragraph.add_run('%s' % l)
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(28)
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            #add cover info
            paragraph = self.document.add_paragraph()
            run_paragraph = paragraph.add_run('Open Source Intelligence Report\n\n\n\n\n\n\n\n\n\n\n')
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(26)
            font.color.rgb = RGBColor(0xe9, 0x58, 0x23)

            paragraph = self.document.add_paragraph()
            run_paragraph = paragraph.add_run('Generated on: %s' % self.today)
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(16)
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            #page break for cover page
            self.document.add_page_break()

            #add intro text on intropage
            heading = self.document.add_heading()
            run_heading = heading.add_run('Executive Summary')
            font = run_heading.font
            font.name = 'Arial'
            font.size = Pt(20)
            font.color.rgb = RGBColor(0xe9, 0x58, 0x23)

            paragraph = self.document.add_paragraph()
            run_paragraph = paragraph.add_run('\nThis document contains information about network, technology, and people associated with the assessment targets. The information was obtained by programatically querying various free or low cost Internet data sources.\n')
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(11)
            run_paragraph = paragraph.add_run('\nThese data include information about the network, technology, and people associated with the targets.\n')
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(11)
            run_paragraph = paragraph.add_run('\nSpecific data sources include: whois, domain name system (DNS) records, Google dork results, and data from recent compromises such as LinkedIn. Other sources include results from Shodan, document metadata from theHarvester and pyFoca, as well as queries to Pastebin, Github, job boards, etc. \n')
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(11)

            #page break for cover page
            self.document.add_page_break()

            heading = self.document.add_heading()
            run_heading = heading.add_run('Table of Contents')
            font = run_heading.font
            font.bold = True
            font.name = 'Arial'
            font.size = Pt(20)
            font.color.rgb = RGBColor(0x0, 0x0, 0x0)

            #TOC https://github.com/python-openxml/python-docx/issues/36
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            font.name = 'Arial'
            font.size = Pt(11)
            fldChar = OxmlElement('w:fldChar')  # creates a new element
            fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = 'TOC \o "1-3" \h \z \u'   # change 1-3 depending on heading levels you need

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:t')
            fldChar3.text = "Right-click to update field."
            fldChar2.append(fldChar3)

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar4)
            p_element = paragraph._p

            #page break for toc
            self.document.add_page_break()

            if cred_result is not None:
                print('[+] Adding credential dump results to report')
                #header
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Credentials found from recent compromises (LinkedIn, Adobe, etc.) related to: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
                paragraph = self.document.add_paragraph()
                for c in cred_result:
                    run_paragraph = paragraph.add_run(''.join(c))
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                self.document.add_page_break()

            #add whois data with header and break after end
            if whois_result is not None:
                print('[+] Adding whois results to report')
                #header
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Whois Data for: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
                #content
                paragraph = self.document.add_paragraph()
                for line in whois_result:
                    if ':' in line:
                        run_paragraph = paragraph.add_run(''.join(line)+'\n')
                        font = run_paragraph.font
                        font.name = 'Arial'
                        font.size = Pt(10)
                self.document.add_page_break()

            #add dns data with header and break after end
            if dns_result is not None:
                print('[+] Adding DNS results to report')
                #header
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Domain Name System Data for: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
                #content
                paragraph = self.document.add_paragraph()
                for d in dns_result:
                    run_paragraph = paragraph.add_run('\n'.join(d))
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(10)
                self.document.add_page_break()

            #google dork output
            if google_result is not None:
                print('[+] Adding google dork results to report')
                #header
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Google Dork Results for: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
                #content
                paragraph = self.document.add_paragraph()
                for r in google_result:
                    run_paragraph = paragraph.add_run(''.join(r+'\n'))
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(10)
                self.document.add_page_break()

            #harvester output
            if theharvester_result is not None:
                print('[+] Adding theHarvester results to report')
                #header
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('theHarvester Results for: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
                #content
                paragraph = self.document.add_paragraph()
                for h in theharvester_result:
                    run_paragraph = paragraph.add_run(''.join(h))
                    #set font stuff
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(10)
                self.document.add_page_break()

            #pastebin scrape output
            if paste_scrape_result is not None:
                print('[+] Adding pastebin scrape results to report')
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Pastebin URLs for %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)

                paragraph = self.document.add_paragraph()
                self.document.add_paragraph(paste_scrape_result)
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(10)
                self.document.add_page_break()

            #general scrape output
            if webscrape_result is not None:
                print('[+] Adding website scraping results to report')
                #header
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Website Scraping Results for %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
                #content
                paragraph = self.document.add_paragraph()
                for sr in webscrape_result:
                    for line in sr:
                        run_paragraph = paragraph.add_run(line)
                        font = run_paragraph.font
                        font.name = 'Arial'
                        font.size = Pt(10)

                self.document.add_page_break()

            #pyfoca results
            if pyfoca_result is not None:
                print('[+] Adding pyfoca results to report')
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('pyFoca Results for: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)

                paragraph = self.document.add_paragraph()
                for fr in pyfoca_result:
                    run_paragraph = paragraph.add_run(''.join(str(fr).strip(("\\ba\x00b\n\rc\fd\xc3"))))
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(10)

                self.document.add_page_break()

            #shodan output
            if shodan_result is not None:
                heading = self.document.add_heading(level=3)
                run_heading = heading.add_run('Shodan Results for: %s' % l)
                font = run_heading.font
                font.name = 'Arial'
                font.color.rgb = RGBColor(0xe9, 0x58, 0x23)

                paragraph = self.document.add_paragraph()
                for shr in shodan_result:
                    try:
                        run_paragraph = paragraph.add_run(str(shr).strip(("\\ba\x00b\n\rc\fd\xc3")))
                        #set font stuff
                        font = run_paragraph.font
                        font.name = 'Arial'
                        font.size = Pt(10)
                    except:
                        print ('probably an encoding error...')
                        continue

            print('[+] Writing file: ./reports/{}/OSINT_{}_.docx'.format(l, l))

            self.document.save(report_directory+l+'/'+l+'OSINT_%s_.docx' % l)
