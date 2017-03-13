#!/usr/bin/env python

#By @arbitrary_code

#Special thanks to:
#@Beamr
#@tatanus
#unum alces!

try:
	import sys
	import time
	import argparse
	import subprocess
	import socket
	import urllib
	import urllib2
	import shodan
	import docx
	from docx.shared import Pt
	from docx.shared import RGBColor
	from docx.shared import Inches
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	from docx.oxml.shared import OxmlElement, qn

	import re
	import os
	from google import search
	import json
	import pprint
	from lxml import html
	import requests
	from collections import Counter

	from webscrape import Scraper
	from whois import Whois
	from dnsquery import Dnsquery
	from hibp import Haveibeenpwned

except ImportError as e:
	raise ImportError('Error importing %s' % e)
	sys.exit(1)



#python-docx: https://pypi.python.org/pypi/python-docx
#shodan: https://github.com/achillean/shodan-python
#google: https://pypi.python.org/pypi/google, also installs beautifulsoup

#*******************************************************************************

def main():
	startTime=time.time()

	#parse input, nargs allows one or more to be entered
	#https://docs.python.org/3/library/argparse.html
	#set nargs back to + for multi search of domain or ip (still really buggy)
	parser = argparse.ArgumentParser()
	parser.add_argument('-a', '--all', help = 'run All queries', action = 'store_true')
	parser.add_argument('-b', '--hibp', help='Search haveibeenpwned.com for breaches related to a domain', action='store_true')
	parser.add_argument('-c', '--creds', help = 'Search local copies of credential dumps', action = 'store_true')
	parser.add_argument('-d', '--domain', nargs = 1, help = 'the Domain you want to search.')
	parser.add_argument('-f', '--foca', help = 'invoke pyfoca', action = 'store_true')
	parser.add_argument('-g', '--googledork', nargs = '+',help = 'query Google for supplied args that are treated as a dork. i.e. -g password becomes a search for "password site:<domain>". Combine terms inside of quotes like "site:rapid7.com inurl:aspx" ')
	parser.add_argument('-i', '--ipaddress', nargs = 1, help = 'the IP address you want to search. Must be a valid IP. ')
	parser.add_argument('-n', '--nslookup',help = 'Name query DNS for supplied -d or -i values. Requires a -d or -i value', action = 'store_true')
	parser.add_argument('-p', '--pastebinsearch', nargs = '+', help = 'Search google for <arg> site:pastebin.com. Requires a pro account if you dont want to get blacklisted.')
	parser.add_argument('-s', '--shodan', help = 'query Shodan, API keys stored in ./api_keys/', action='store_true')
	parser.add_argument('-S', '--scraper', help = 'Scrape pastebin, github, indeed, more to be added. API keys stored in ./api_keys/', action = 'store_true')
	parser.add_argument('-t', '--theharvester', help = 'Invoke theHarvester', action = 'store_true')
	parser.add_argument('-v', '--verbose', help = 'Verbose', action = 'store_true')	
	parser.add_argument('-w', '--whois', help = 'query Whois for supplied -d or -i values. Requires a -d or -i value', action = 'store_true')
	
	args = parser.parse_args()

	#verbosity flag to print logo and args
	if args.verbose is True:print '''
    _         _    ___  ____ ___ _   _ _____ 
   / \  _   _| |_ / _ \/ ___|_ _| \ | |_   _|
  / _ \| | | | __| | | \___ \| ||  \| | | |  
 / ___ \ |_| | |_| |_| |___) | || |\  | | |  
/_/   \_\__,_|\__|\___/|____/___|_| \_| |_|\n'''

	if args.verbose is True:print 'AutOSINT.py v0.1, a way to do some automated OSINT tasks\n'
	if args.verbose is True:print args

	#set True on action store_true args if -a
	if args.all is True:
		args.creds = True
		args.hibp = True
		args.foca = True
		args.nslookup = True
		args.theharvester = True
		args.whois = True
		args.scraper = True
		args.shodan = True
		if args.googledork is None:
			print '[-] You need to provide arguments for google dorking. e.g -g inurl:apsx'
			sys.exit()
		if args.pastebinsearch is None:
			print '[-] You need to provide arguments for pastebin keywords. e.g -p password id_rsa'
			sys.exit()

	#check directories
	reportDir='./reports/'
	apiKeyDir='./api_keys/'
	if not os.path.exists(reportDir):
		os.makedirs(reportDir)

	if not os.path.exists(apiKeyDir):
		os.makedirs(apiKeyDir)

	#validate entered IP address? do we even care about IP address? i and d do the same shit
	if args.ipaddress is not None:
		for a in args.ipaddress:
			try:
				socket.inet_aton(a)
			except socket.error:
				print '[-] Invalid IP address entered!' + a
				sys.exit()

	#require at least one argument
	if not (args.domain or args.ipaddress):
	    parser.error('[-] No OSINT reference provided, add domain(s) with -d or IP address(es) with -i\n')
	    sys.exit()

	#if no queries defined, exit. -a sets all so we're good there
	if (args.whois is False and \
		args.hibp is False and \
		args.nslookup is False and \
		args.googledork is None and \
		args.shodan is False and \
		args.creds is False and \
		args.theharvester is False and \
		args.scraper is False and \
		args.pastebinsearch is None and \
		args.foca is False):
		print '[-] No options specified, use -h or --help for a list'
		sys.exit()

	#check to see if an ip or domain name was entered
	if args.domain is not None:
		for d in args.domain:
			lookup = args.domain
			for l in lookup:
				if not os.path.exists(reportDir+'/'+l):
					os.makedirs(reportDir+'/'+l)
				
	else:
		for i in args.ipaddress:
			lookup = args.ipaddress
			for l in lookup:
				if not os.path.exists(reportDir+'/'+l):
					os.makedirs(reportDir+'/'+l)

	if args.verbose is True:
		print '[+] Lookup Values: '+', '.join(lookup)


	#init results lists
	
	whoisResult=[]
	dnsResult = []
	googleResult =[]
	shodanResult=[]
	pasteScrapeResult = []
	pasteScrapeContent = []
	harvesterResult =[]
	scrapeResult=[]
	credResult=[]
	pyfocaResult=[]
	hibpResult=[]

	#call function if -w arg
	if args.whois is True:
		
		whoisSearch = Whois()
		whoisSearch.run(args, lookup, reportDir)
		#whoisResult = whois_search(args, lookup, reportDir)


	#call function if -n arg
	if args.nslookup is True:
		
		dnsQuery = Dnsquery()
		dnsQuery.run(args, lookup, reportDir)
		#dnsResult = dns_search(args, lookup, reportDir)

	#call function if -b arg
	if args.hibp is True:

		hibpSearch = Haveibeenpwned()
		hibpSearch.run(args, lookup, reportDir)
		#hibpResult = hibp_search(args, lookup, reportDir)
	

	#call function if -g arg
	if args.googledork is not None:
		googleResult=google_search(args, lookup, reportDir)

	#call function if -s arg
	if args.shodan is True:
		shodanResult = shodan_search(args, lookup, reportDir, apiKeyDir)
	
	#call function if -p arg
	if args.pastebinsearch is not None:
		pasteScrapeResult=pastebin_search(args, lookup, reportDir, apiKeyDir)
	
	# call function if -t arg
	if args.theharvester is True:
		harvesterResult=the_harvester(args, lookup, reportDir)
	
	#call function if -c arg 
	if args.creds is True:
		credResult=credential_leaks(args, lookup, startTime, reportDir)
	

	#call function if -S arg
	if args.scraper is True:
		
		web_scraper=Scraper()

		scrapeResults = web_scraper.run(args, lookup, reportDir, apiKeyDir)

	#call function if -f arg
	if args.foca is True:
		pyfocaResult=pyfoca(args, lookup, reportDir)

	#if args.hibp is True:
		#hibpResult=hibp_search(args, lookup, reportDir, apiKeyDir)


	#run the docx report. text files happen in the respective functions
	write_report(args, reportDir, lookup, whoisResult, dnsResult, googleResult, shodanResult, pasteScrapeResult, harvesterResult, scrapeResult, credResult, pyfocaResult)


#*******************************************************************************
#ssl scan 
#*******************************************************************************
#censys
#https://www.censys.io/ipv4?q=rapid7.com
#rest api  https://www.censys.io/api/v1/
#*******************************************************************************
#https://dnsdumpster.com/
#cool mapping of AS, etc
#*******************************************************************************
#passive dns

#*******************************************************************************
#viewdns.info
#http://viewdns.info/api/
#*******************************************************************************
#he bgp info
#http://bgp.he.net/dns/rapid7.com#_ipinfo
#*******************************************************************************
#active osint:
#zone transfer host -a does this?
#ike endpoints
#http screnshots
#*******************************************************************************
#salesforce api
#*******************************************************************************
#recon-ng

#*******************************************************************************
# this could be GREATLY improved. need to implement the custom search api
# pass google dorks as args for now
# GHDB password dorks https://www.exploit-db.com/google-hacking-database/9/
# GHDB sensitive dirs https://www.exploit-db.com/google-hacking-database/3/
# uses this awesome module https://pypi.python.org/pypi/google
# requires beautifulsoup
#https://stackoverflow.com/questions/4082966/what-are-the-alternatives-now-that-the-google-web-search-api-has-been-deprecated/11206266#11206266

def google_search(args, lookup, reportDir):
	#need a default dork list

	#C58EA28C-18C0-4a97-9AF2-036E93DDAFB3 is string for open OWA attachments

	#init lists
	googleResult = []

	#if no args provided, default to 'password'. need an inner function?
	if args.googledork is None:
		print '[i] No dork args, defaulting to "password"'
		args.googledork = ['password']

	else:

		for d in args.googledork:
			if args.verbose is True:print '[+] Google dorking for: %s' % d
			
			#default to password if no arg

			#iterate the lookup list
			for i, l in enumerate(lookup):
				googleResult.append('Google query for: '+str(d)+ ' ' + 'site:'+str(l))
				googleFile=open(reportDir+l+'/'+l+'_google_dork_'+str(d)+'.txt','w')

				#show user whiat is being searched
				print '[+] Google query ' + str(i + 1) + ' for "'+str(d)+' ' + 'site:'+str(l) + '"'
				
				try:
					#iterate url results from search of password(for now) and site:current list value
					for url in search(str(d)+ ' ' + 'site:'+str(l), stop = 20):
						
						#append results together
						googleResult.append(url)

						#rate limit to 1 per second
						time.sleep(1)
				#catch exceptions
				except Exception:
					pass
		#iterate results
		for r in googleResult:
			#write results on newlines
			googleFile.writelines(r + '\r\n')

		#verbosity flag
		if args.verbose is True:
			for r in googleResult: print ''.join(r)
				
		#return results list
		return googleResult
	
		

#*******************************************************************************
def shodan_search(args, lookup, reportDir, apiKeyDir):
	#probably need to customize search type based on -i or -d		
	#first if  https://shodan.readthedocs.io/en/latest/tutorial.html#connect-to-the-api
	#else https://shodan.readthedocs.io/en/latest/tutorial.html#looking-up-a-host

	#list that we'll return
	shodanResult = []
	shodanApiKey=''

	#check for api key file
	if not os.path.exists(apiKeyDir + 'shodan.key'):
		print '[-] You are missing %s/shodan.key' % apiKeyDir
		#shodanApiKey=raw_input("Please provide an API Key: ")

	#read API key
	try:
		with open(apiKeyDir + 'shodan.key', 'r') as apiKeyFile:
			for k in apiKeyFile:
				shodanApiKey = k
	except:
		print '[-] Error opening %s/shodan.key key file, skipping. ' % apiKeyDir

	#invoke api with api key provided
	shodanApi = shodan.Shodan(shodanApiKey)

	#roll through the lookup list from -i or -d
	for i, l in enumerate(lookup):
		#open output file
		shodanFile=open(reportDir+l+'/'+l+'_shodan.txt','w')



		#maybe just do the rest here instead of py api client?
		#url='https://exploits.shodan.io/api/search?query='+str(l)+'&key='+str(shodanApiKey)


		#user notification that something is happening
		print '[+] Querying Shodan via API search for ' + l
		try:
			#set results to api search of current lookup value
			#https://shodan.readthedocs.io/en/latest/examples/basic-search.html
			result = shodanApi.search(query="hostname:"+l)
			print '[+] Shodan found: '+str(result['total'])+' hosts'
			#for each result
			for service in result['matches']:
				if args.verbose is True:print str(service['ip_str'].encode('utf-8')+\
					' ISP: '+service['isp'].encode('utf-8')+\
					' Last seen: '+service['timestamp'].encode('utf-8'))
				if args.verbose is True:print service['data'].encode('utf-8')

				#append to shodanResult list
				shodanResult.append(str(\
					service['ip_str'].encode('utf-8')+\
					'\nISP:'+service['isp'].encode('utf-8')+\
					'\nLast seen:'+service['timestamp'].encode('utf-8'))+\
					'\n'+service['data'].encode('utf-8'))				



			'''exploits=shodanApi.exploits.search(l)
			for ex in exploits:
				shodanResult.append(str(ex))'''
		#catch exceptions		
		except shodan.APIError, e:
			#print excepted error
			print '[-] Shodan Error: %s' % e + ' Skipping!!!'
			print '[!] You may need to specify an API key with -s <api key>'
			return
			
	#write contents of shodanResult list. this needs formatted
	shodanFile.writelines('[+] Shodan found: '+str(result['total'])+' hosts\n\n')
	shodanFile.writelines(shodanResult)

	return shodanResult


	
#*******************************************************************************
#right now this just google dorks a supplied arg for site:pastebin.com
#need to implement scraping api http://pastebin.com/api_scraping_faq
#scraping url is here http://pastebin.com/api_scraping.php
def pastebin_search(args, lookup, reportDir, apiKeyDir):
	
	userAgent = {'User-agent': 'Mozilla/5.0'}
	
	#return values
	pasteScrapeUrl = []
	pasteScrapeContent = []
	pasteScrapeResult =[]

	# check for empty args
	if args.pastebinsearch is not None:


		for a in args.pastebinsearch:
			#init lists
			scrapeURL = []
			scrapeContent = []

			#iterate the lookup list
			for i, l in enumerate(lookup):

				#init textfiles
				scrapedFile=open(reportDir+l+'/'+l+'_pastebin_content.txt','w')
				pasteUrlFile=open(reportDir+l+'/'+l+'_pastebin_urls.txt','w')
				
				#show user whiat is being searched
				print '[+] Searching Pastebin for public pastes containing %s' % (l)
				print '[i] May require a Pastebin Pro account for IP whitelisting'


				#run google query code
				try:
					#iterate url results from search of dork arg and supplied lookup value against pastebin. return top 20 hits
					for url in search(str(a) +' '+ str(l) + ' site:pastebin.com', stop = 20):
						#delay 1 second to be polite
						time.sleep(1)
						#append results together
						scrapeURL.append(url)
						if args.verbose is True:print '[+] Paste containing "%s" and "%s" found at: %s' (a,l,url)
				except Exception:
					print '[-] Error dorking pastebin URLs, skipping...'
					pasteScrapeResult.append('Error scraping Pastebin')
					continue

				for u in scrapeURL:
					#http://docs.python-guide.org/en/latest/scenarios/scrape/
					try:
						page = requests.get(u, headers = userAgent)
						pasteUrlFile.writelines(u)
					except:
						print '[-] Error opening ' + u +':'
						pasteScrapeResult.append('Error opening %s' % u)
						continue


					#build html tree
					tree = html.fromstring(page.content)

					#if verbose spit out url, search term and domain searched
					if args.verbose is True:print '[+] Looking for instances of %s and %s in %s \n' % (a,l,url)
					#grab raw paste data from the textarea
					rawPasteData = tree.xpath('//textarea[@class="paste_code"]/text()')

					#search lines for lookup and keyword
					for line in rawPasteData:
						#regex for the lookup value (domain) in that line
						#if re.search((str(l)), line):
						if str(l) in line:
							#if the argument search term is in the line
							if a in line:
								scrapedFile.writelines(a)

				return pasteScrapeResult

				
#*******************************************************************************
def the_harvester(args, lookup, reportDir):

	#https://github.com/laramies/theHarvester
	if args.theharvester is True:
		
		#init lists
		harvested = []
		harvesterGoogleResult = []
		harvesterLinkedinResult = []
		harvesterResult=[]

		#based on domain or ip, enumerate with index and value
		for i, l in enumerate(lookup):

			#open file to write to
			harvesterFile=open(reportDir+l+'/'+l+'_theharvester.txt','w')

			#run harvester with -b google on lookup
			try:
				print '[+] Running theHarvester -b google -d %s ' % l
				harvesterGoogleCmd = subprocess.Popen(['theharvester', '-b', 'google', '-d', str(l), '-l', '500', '-h'], stdout = subprocess.PIPE).communicate()[0].split('\r\n')
			except:
				print '[-] Error running theharvester. Make sure it is in your PATH and you are connected to the Internet'
				harvesterResult.append('Error running theHarvester')
				continue

			#run harvester with -b linkedin on lookup
			try:
				print '[+] Running theHarvester -b linkedin -d %s ' % l
				harvesterLinkedinCmd = subprocess.Popen(['theharvester', '-b', 'linkedin', '-d', str(l), '-l', '500', '-h'], stdout = subprocess.PIPE).communicate()[0].split('\r\n')
			except:
				print '[-] Error running theharvester. Make sure it is in your PATH and you are connected to the Internet'
				harvesterResult.append('Error running theHarvester')
				continue

			#append lists together
			harvesterResult.append(harvesterGoogleCmd)
			harvesterResult.append(harvesterLinkedinCmd)

			#append resutls and write to lookup result file
			for r in harvesterResult:
				harvesterFile.writelines(r)

				
		#verbosity
		if args.verbose is True:
			for h in harvesterResult: print '\n'.join(h)


		#return list object
		return harvesterResult




#*******************************************************************************
def credential_leaks(args, lookup, startTime, reportDir):
	#grep through local copies of various password database dumps. 
	#compares to a hashcat potfile as well
	#you'll need a ./credleaks directory and a ./potfile directory populated
	#dumps need to be in uname:hash format
	#this could probably stand to be multi threaded

	potfileDir = './potfile'
	credLeakDir = './credleaks'

	if args.creds is True:

		if not os.path.exists(potfileDir):
			print '[-] The potfile directory is missing. Symlink your location to ./potfile and see if that works'
			return
		

		if not os.path.exists(credLeakDir):
			print '[-] The credleaks directory is missing. Symlink your location to ./credleaks and see if that works'
			return


	
		#for each domain/ip provided
		for l in lookup:
			credFile=open(reportDir+l+'/'+l+'_creds.txt','w')

			#init dictionary
			dumpDict={}
			credResult=[]


			print '[+] Searching credential dumps for entries that contain '+l
			#overall, take the lookup value (preferably a domain) and search the dumps for it
			#for each file in ./credleaks directory
			#really need to get this data out of text files an into an indexed form. it's slow af 
			for credFileName in os.listdir('./credleaks/'):
				#open the file
				credFileOpen = open('./credleaks/'+credFileName, "r")
				j=0
				#i=0
				#for each line in opened file
				for line in credFileOpen:
					#line counter index. i thought maybe i could also display how many lines were searched
					#i=i+1
					#regex search for our current lookup value l
					#if re.search((str(l)), line):
					if str(l) in line:
						#counter index
						j=j+1
						#look for a colon delimiter. dump files should be like email:hash. this of course assumes the creds file has emails as usernames
						if ':' in line:
							#split matches based on colons, sorta like 'awk -F :'. emails shouldnt have colons, right?
							#also the dat HAS to require colons otherwise it will return an index error
							matchedLine=line.split(":")
							#take the split parts, 0 and 1 that are uname and hash, respectively
							#place into a dict and strip the \r\n off of them
							dumpDict[str(matchedLine[1].rstrip("\r\n"))]=str(matchedLine[0].rstrip("\r\n"))
						#otherwise print xxx if theres no hash for the entry. some dumps dont have hashes for everyone...
						else:
							dumpDict['xxx']=str(line.rstrip("\r\n"))
				#print each file searched and how many matches if verbose
				if args.verbose is True: 
					print '[i] Searched ' + str(credFileName)+' and found '+ str(j)

			
			#print hash and user of files if verbose	
			if args.verbose is True:
				for h, u in dumpDict.items():
					print(str(u)) 

			#start printing stuff and appending to credResult
			print '[+] Searching Local Credential Dumps in ./credleaks against potfile in ./potfile '
			credFile.writelines('********EMAILS FOUND BELOW********\n\n\n\n')
			credResult.append('********EMAILS FOUND BELOW********\n\n\n\n')
			
			#iterate the dictionary containing user and hashes
			for h, u in dumpDict.items():
				#write username to text file
				credFile.writelines(str(u)+'\n')
				#write username to credResult for the docx report
				credResult.append(str(u)+'\n')
				
			credFile.writelines('********CREDENTIALS FOUND BELOW*********\n\n\n\n')
			credResult.append('********CREDENTIALS FOUND BELOW*********\n\n\n\n')
			
			#this section 'cracks' the hashes provided a pre-populated pot file
			#still in our lookup value iterate potfiles directory. you can have multiple pots, just in case
			for potFileName in os.listdir('./potfile/'):
				#open a pot file
				with open('./potfile/'+potFileName, 'r') as potFile:
					#tell user you are looking
					print '[i] Any creds you have in your potfile will appear below as user:hash:plain : '
					#then look at every line
					for potLine in potFile:
						#then for every line look at every hash and user in the dict
						for h, u in dumpDict.items():
							#if the hash in the dict matches a line in the potfile
							#that is also the same length as the original hash (this is probably a crappy check tho...)
							if str(h) == str(potLine[0:len(h)]):
								#print the user: and the line from the potfile (hash:plain) to the user
								print str(u)+':'+str(potLine.rstrip("\r\n"))
								#need to append the output to a variable to return or write to the file
								#this is separate because not all found usernames/emails have hashes and not all hashes are cracked
								#write to text file
								credFile.writelines(str(u)+':'+str(potLine[len(h):]))
								#add to credResult for docx report
								credResult.append(str(u)+':'+str(potLine[len(h):]))


			return credResult	
			print credResult

#*******************************************************************************
def pyfoca(args, lookup, reportDir):
	#https://github.com/altjx/ipwn
	if args.foca is True:
		
		#init lists
		pyfocaResult=[]

		#based on domain or ip, enumerate with index and value
		for i, l in enumerate(lookup):

			#open file to write to
			pyfocaFile=open(reportDir+l+'/'+l+'_pyfoca.txt','w')

			#run pyfoca with -d domain. should automagically do metadata
			try:
				print '[+] Running pyfoca -d %s' % l
				pyfocaCmd = subprocess.Popen(['pyfoca', '-d', str(l)], stdout = subprocess.PIPE).communicate()[0].split('\r\n')
			except:
				print '[-] Error running pyfoca. Make sure it is in your PATH and you are connected to the Internet'
				pyfocaResult.append('Error running pyfoca')
				pyfocaFile.writelines('Error running pyfoca')
				continue
			
			#append output
			pyfocaFile.writelines(pyfocaCmd)
			pyfocaResult.append(pyfocaCmd)
			
			#spew if verbose
			if args.verbose is True: 
				for p in pyfocaResult:print '\n'.join(p)

			return pyfocaResult

#*******************************************************************************

def write_report(args, reportDir, lookup, whoisResult, dnsResult, googleResult, shodanResult, pasteScrapeResult, harvesterResult, scrapeResult, credResult, pyfocaResult):

	today = time.strftime("%m/%d/%Y")
	for l in lookup:
		print '[+] Starting OSINT report for '+l

		#dump to a word doc
		#refs
		#https://python-docx.readthedocs.io/en/latest/user/text.html
		#https://python-docx.readthedocs.io/en/latest/user/quickstart.html
		
		#create a document 
		document = docx.Document()


		#add logo
		document.add_picture('./resources/logo.png', height=Inches(1.25))

		#add domain cover info

		paragraph = document.add_paragraph() 
		runParagraph = paragraph.add_run('%s' % l)
		font=runParagraph.font
		font.name = 'Arial'
		font.size = Pt(28)
		font.color.rgb = RGBColor(0x00,0x00,0x00)
	
		#add cover info
		paragraph = document.add_paragraph() 
		runParagraph = paragraph.add_run('Open Source Intelligence Report\n\n\n\n\n\n\n\n\n\n\n')
		font=runParagraph.font
		font.name = 'Arial'
		font.size = Pt(26)
		font.color.rgb = RGBColor(0xe9,0x58,0x23)

		paragraph = document.add_paragraph() 
		runParagraph = paragraph.add_run('Generated on: %s' % today)
		font=runParagraph.font
		font.name = 'Arial'
		font.size = Pt(16)
		font.color.rgb = RGBColor(0x00,0x00,0x00)


		#page break for cover page
		document.add_page_break()
		
		#add intro text on intropage

		heading = document.add_heading()
		runHeading = heading.add_run('Executive Summary')
		font=runHeading.font
		font.name = 'Arial'
		font.size = Pt(20)
		font.color.rgb = RGBColor(0xe9,0x58,0x23)

		paragraph = document.add_paragraph() 
		runParagraph = paragraph.add_run('\nThis document contains information about network, technology, and people associated with the assessment targets. The information was obtained by programatically querying various free or low cost Internet data sources.\n')
		font=runParagraph.font
		font.name = 'Arial'
		font.size = Pt(11)
		runParagraph = paragraph.add_run('\nThese data include information about the network, technology, and people associated with the targets.\n')
		font=runParagraph.font
		font.name = 'Arial'
		font.size = Pt(11)
		runParagraph = paragraph.add_run('\nSpecific data sources include: whois, domain name system (DNS) records, Google dork results, and data from recent compromises such as LinkedIn. Other sources include results from Shodan, document metadata from theHarvester and pyFoca, as well as queries to Pastebin, Github, job boards, etc. \n')
		font=runParagraph.font
		font.name = 'Arial'
		font.size = Pt(11)

		
		#page break for cover page
		document.add_page_break()

		heading = document.add_heading()
		runHeading = heading.add_run('Table of Contents')
		font=runHeading.font
		font.bold = True
		font.name = 'Arial'
		font.size = Pt(20)
		font.color.rgb = RGBColor(0x0,0x0,0x0)

		#TOC https://github.com/python-openxml/python-docx/issues/36
		paragraph = document.add_paragraph()
		run = paragraph.add_run()
		font.name = 'Arial'
		font.size = Pt(11)
		fldChar = OxmlElement('w:fldChar')  # creates a new element
		fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

		instrText = OxmlElement('w:instrText')
		instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
		instrText.text = 'TOC \o "1-3" \h \z \u'   # change 1-3 depending on heading levels you need

		fldChar2 = OxmlElement('w:fldChar')
		fldChar2.set(qn('w:fldCharType'), 'separate')
		fldChar3 = OxmlElement('w:t')
		fldChar3.text = "Right-click to update field."
		fldChar2.append(fldChar3)

		fldChar4 = OxmlElement('w:fldChar')
		fldChar4.set(qn('w:fldCharType'), 'end')

		r_element = run._r
		r_element.append(fldChar)
		r_element.append(instrText)
		r_element.append(fldChar2)
		r_element.append(fldChar4)
		p_element = paragraph._p



		#page break for toc
		document.add_page_break()


		if credResult:
			print '[+] Adding credential dump results to report'
			#header
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('Credentials found from recent compromises (LinkedIn, Adobe, etc.) related to: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)
			paragraph = document.add_paragraph()
			for c in credResult:
				runParagraph = paragraph.add_run(''.join(c))
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(11)
			document.add_page_break()
		
		#add whois data with header and break after end
		if whoisResult:
			print '[+] Adding whois results to report'
			#header
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('Whois Data for: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)
			#content
			paragraph = document.add_paragraph()
			for w in whoisResult:
				runParagraph = paragraph.add_run('\n'.join(w))
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(10)
			document.add_page_break()
		
		#add dns data with header and break after end
		if dnsResult:
			print '[+] Adding DNS results to report'
			#header
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('Domain Name System Data for: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)
			#content
			paragraph = document.add_paragraph()
			for d in dnsResult:
				runParagraph = paragraph.add_run('\n'.join(d))
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(10)
			document.add_page_break()

		#google dork output
		if googleResult:
			print '[+] Adding google dork results to report'
			#header
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('Google Dork Results for: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)
			#content
			paragraph = document.add_paragraph()
			for r in googleResult:
				runParagraph = paragraph.add_run(''.join(r+'\n'))
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(10)
			document.add_page_break()
		
		#harvester output
		if harvesterResult:
			print '[+] Adding theHarvester results to report'
			#header
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('theHarvester Results for: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)
			#content
			paragraph = document.add_paragraph()
			for h in harvesterResult: 
				runParagraph = paragraph.add_run(''.join(h))
				#set font stuff
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(10)
			document.add_page_break()
		

		#pastebin scrape output
		if pasteScrapeResult:
			print '[+] Adding pastebin scrape results to report'
			document.add_heading('Pastebin URLs for %s' % l, level=3)
			document.add_paragraph(pasteScrapeResult)
			document.add_page_break()
			#document.add_paragraph(pasteScrapeContent)
			#document.add_page_break()



		#general scrape output
		if scrapeResult:
			print '[+] Adding website scraping results to report'
			#header
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('Website Scraping Results for %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)
			#content
			paragraph = document.add_paragraph()
			for sr in scrapeResult:
				runParagraph = paragraph.add_run(sr)
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(10)

			document.add_page_break()


		#pyfoca results
		if pyfocaResult:
			print '[+] Adding pyfoca results to report'
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('pyFoca Results for: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)


			paragraph = document.add_paragraph()
			for fr in pyfocaResult:
				#lolwut
				runParagraph = paragraph.add_run(''.join(str(fr).strip(("\\ba\x00b\n\rc\fd\xc3"))))
				font=runParagraph.font
				font.name = 'Arial'
				font.size = Pt(10)

			document.add_page_break()
		
		#shodan output
		if shodanResult:
			heading = document.add_heading(level=3)
			runHeading = heading.add_run('Shodan Results for: %s' % l)
			font=runHeading.font
			font.name = 'Arial'
			font.color.rgb = RGBColor(0xe9,0x58,0x23)


			paragraph = document.add_paragraph()
			for shr in shodanResult:
				try:
					runParagraph = paragraph.add_run(str(shr).strip(("\\ba\x00b\n\rc\fd\xc3")))
					#set font stuff
					font=runParagraph.font
					font.name = 'Arial'
					font.size = Pt(10)
				except:
					print 'probably an encoding error...'
					continue
		
		print '[+] Writing file: ./reports/%s/OSINT_%s_.docx'  % (l, l)
		document.save(reportDir+l+'/'+l+'OSINT_%s_.docx' % l)


if __name__ == '__main__':
    main()
